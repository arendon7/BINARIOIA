import tempfile, unittest
from pathlib import Path
from unittest.mock import patch
from docx import Document
from documentos_ia.exporters import export_pdf_from_docx

class PDFR7Tests(unittest.TestCase):
    def test_pdf_fallback_without_libreoffice(self):
        with tempfile.TemporaryDirectory() as td:
            td=Path(td);docx=td/'sample.docx';d=Document();d.add_heading('Prueba',0);d.add_paragraph('PDF interno sin LibreOffice.');d.save(docx)
            with patch('documentos_ia.exporters.shutil.which',lambda name: None), patch('documentos_ia.exporters.Path.exists',Path.exists):
                pdf=export_pdf_from_docx(docx,td/'out')
            self.assertTrue(pdf.exists());self.assertGreater(pdf.stat().st_size,500)

if __name__=='__main__':unittest.main()
