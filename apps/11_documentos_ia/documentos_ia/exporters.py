from __future__ import annotations
from pathlib import Path
import html, subprocess, shutil
from .models import DocumentSpec
from .quality import evaluate


def export_markdown(spec: DocumentSpec, out: str|Path, require_pass=True) -> Path:
    if require_pass and evaluate(spec).status!='pass':
        raise ValueError('Quality Gate bloqueó la exportación final.')
    lines=[f'# {spec.title}','']
    for b in spec.blocks:
        if b.kind=='heading': lines += [f"{'#'*(min(5,b.level)+1)} {b.text}",'']
        elif b.kind=='paragraph': lines += [b.text,'']
        elif b.kind=='bullets': lines += [*(f'- {x}' for x in b.items),'']
        elif b.kind=='numbered': lines += [*(f'{i}. {x}' for i,x in enumerate(b.items,1)),'']
        elif b.kind=='quote': lines += [f'> {b.text}','']
        elif b.kind=='callout': lines += [f'**{b.text}**','']
        elif b.kind=='table' and b.rows:
            rows=b.rows
            lines += ['| '+' | '.join(rows[0])+' |','| '+' | '.join('---' for _ in rows[0])+' |']
            lines += ['| '+' | '.join(r)+' |' for r in rows[1:]]+['']
        elif b.kind=='page_break': lines += ['<div style="page-break-after:always"></div>','']
    p=Path(out); p.write_text('\n'.join(lines),encoding='utf-8'); return p


def export_html(spec: DocumentSpec, out: str|Path, require_pass=True) -> Path:
    if require_pass and evaluate(spec).status!='pass':
        raise ValueError('Quality Gate bloqueó la exportación final.')
    css='''body{font-family:Arial,sans-serif;max-width:900px;margin:42px auto;color:#172033;line-height:1.55;padding:0 24px}h1{text-align:center;color:#17365D;margin-bottom:36px}h2,h3,h4{color:#17365D;margin-top:28px}p{text-align:justify}table{width:100%;border-collapse:collapse;margin:18px 0}th,td{border:1px solid #ccd5df;padding:8px}blockquote{border-left:4px solid #2F75B5;padding-left:14px;color:#45566c}.sources{margin-top:35px;font-size:12px;color:#667}'''
    body=[f'<h1>{html.escape(spec.title)}</h1>']
    for b in spec.blocks:
        if b.kind=='heading': body.append(f'<h{min(6,b.level+1)}>{html.escape(b.text)}</h{min(6,b.level+1)}>')
        elif b.kind=='paragraph': body.append(f'<p>{html.escape(b.text)}</p>')
        elif b.kind in {'bullets','numbered'}:
            tag='ul' if b.kind=='bullets' else 'ol'; body.append('<'+tag+'>'+''.join(f'<li>{html.escape(x)}</li>' for x in b.items)+'</'+tag+'>')
        elif b.kind=='quote': body.append(f'<blockquote>{html.escape(b.text)}</blockquote>')
        elif b.kind=='callout': body.append(f'<p><strong>{html.escape(b.text)}</strong></p>')
        elif b.kind=='table' and b.rows:
            body.append('<table><thead><tr>'+''.join(f'<th>{html.escape(x)}</th>' for x in b.rows[0])+'</tr></thead><tbody>'+''.join('<tr>'+''.join(f'<td>{html.escape(x)}</td>' for x in r)+'</tr>' for r in b.rows[1:])+'</tbody></table>')
        elif b.kind=='page_break': body.append('<div style="page-break-after:always"></div>')
    if spec.sources:
        body.append('<div class="sources"><h2>Fuentes y procedencia</h2><ol>'+''.join(f'<li>{html.escape(s.title)} - SHA-256 {html.escape(s.hash_sha256[:16])}...</li>' for s in spec.sources)+'</ol></div>')
    doc=f'<!doctype html><html lang="{spec.language}"><meta charset="utf-8"><title>{html.escape(spec.title)}</title><style>{css}</style><body>{"".join(body)}</body></html>'
    p=Path(out); p.write_text(doc,encoding='utf-8'); return p


def _set_cell_shading(cell, fill='D9EAF7'):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    tcPr=cell._tc.get_or_add_tcPr(); shd=OxmlElement('w:shd'); shd.set(qn('w:fill'),fill); tcPr.append(shd)


def _add_page_number(paragraph):
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run=paragraph.add_run(); fldChar1=OxmlElement('w:fldChar'); fldChar1.set(qn('w:fldCharType'),'begin')
    instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=' PAGE '
    fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'end')
    run._r.extend([fldChar1,instr,fldChar2])


def export_docx(spec: DocumentSpec, out: str|Path, require_pass=True) -> Path:
    if require_pass and evaluate(spec).status!='pass':
        raise ValueError('Quality Gate bloqueó la exportación final.')
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.section import WD_SECTION
    from docx.enum.table import WD_TABLE_ALIGNMENT
    doc=Document()
    sec=doc.sections[0]
    margin=Cm(spec.style.margins_cm)
    sec.top_margin=margin; sec.bottom_margin=margin; sec.left_margin=margin; sec.right_margin=margin
    styles=doc.styles
    normal=styles['Normal']; normal.font.name=spec.style.font_name; normal.font.size=Pt(spec.style.body_size_pt)
    normal.paragraph_format.line_spacing=spec.style.line_spacing
    for name in ['Title','Heading 1','Heading 2','Heading 3']:
        st=styles[name]; st.font.name=spec.style.font_name
        st.font.color.rgb=RGBColor.from_string(spec.style.heading_color)
    title=doc.add_paragraph(style='Title'); title.add_run(spec.title)
    if spec.style.title_centered: title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    if spec.objective:
        p=doc.add_paragraph(); r=p.add_run('Objetivo: '); r.bold=True; p.add_run(spec.objective); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    for b in spec.blocks:
        if b.kind=='heading': doc.add_heading(b.text,level=max(1,min(3,b.level)))
        elif b.kind=='paragraph':
            p=doc.add_paragraph(b.text); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY if spec.style.body_justified else WD_ALIGN_PARAGRAPH.LEFT
        elif b.kind=='bullets':
            for x in b.items: doc.add_paragraph(x,style='List Bullet')
        elif b.kind=='numbered':
            for x in b.items: doc.add_paragraph(x,style='List Number')
        elif b.kind=='quote':
            p=doc.add_paragraph(); r=p.add_run(b.text); r.italic=True; p.paragraph_format.left_indent=Cm(1)
        elif b.kind=='callout':
            p=doc.add_paragraph(); r=p.add_run(b.text); r.bold=True
        elif b.kind=='table' and b.rows:
            table=doc.add_table(rows=len(b.rows),cols=max(len(r) for r in b.rows)); table.alignment=WD_TABLE_ALIGNMENT.CENTER; table.style='Table Grid'
            for i,row in enumerate(b.rows):
                for j,val in enumerate(row): table.cell(i,j).text=str(val)
            for cell in table.rows[0].cells:
                _set_cell_shading(cell, spec.style.accent_color); 
                for run in cell.paragraphs[0].runs: run.font.color.rgb=RGBColor(255,255,255); run.bold=True
        elif b.kind=='page_break': doc.add_page_break()
    if spec.sources:
        doc.add_heading('Fuentes y procedencia',level=1)
        for s in spec.sources:
            doc.add_paragraph(f'{s.title} - SHA-256 {s.hash_sha256[:16]}...',style='List Number')
    footer=sec.footer.paragraphs[0]; footer.text=spec.style.footer_text; footer.alignment=WD_ALIGN_PARAGRAPH.CENTER
    if spec.style.include_page_numbers:
        footer.add_run(' · pág. '); _add_page_number(footer)
    p=Path(out); doc.save(str(p)); return p


def export_pdf_from_docx(docx_path: str|Path, out_dir: str|Path) -> Path:
    src=Path(docx_path); outdir=Path(out_dir); outdir.mkdir(parents=True,exist_ok=True)
    candidates=[shutil.which('libreoffice'),shutil.which('soffice')]
    if Path('/Applications/LibreOffice.app/Contents/MacOS/soffice').exists(): candidates.append('/Applications/LibreOffice.app/Contents/MacOS/soffice')
    soffice=next((x for x in candidates if x),None)
    pdf=outdir/(src.stem+'.pdf')
    if soffice:
        cp=subprocess.run([soffice,'--headless','--convert-to','pdf','--outdir',str(outdir),str(src)],capture_output=True,text=True,timeout=120)
        if cp.returncode==0 and pdf.exists(): return pdf
    # Fallback interno: PDF usable sin obligar a instalar LibreOffice.
    try:
        from docx import Document
        from reportlab.lib import colors
        from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle
        doc=Document(str(src)); styles=getSampleStyleSheet()
        title=ParagraphStyle('SBTitle',parent=styles['Title'],alignment=TA_CENTER,spaceAfter=18)
        body=ParagraphStyle('SBBody',parent=styles['BodyText'],alignment=TA_JUSTIFY,leading=15,spaceAfter=8)
        story=[]
        for para in doc.paragraphs:
            text=(para.text or '').strip()
            if not text: story.append(Spacer(1,5)); continue
            style_name=(para.style.name or '').lower() if para.style else ''
            if 'title' in style_name: st=title
            elif 'heading 1' in style_name: st=styles['Heading1']
            elif 'heading 2' in style_name: st=styles['Heading2']
            elif 'heading 3' in style_name: st=styles['Heading3']
            else: st=body
            story.append(Paragraph(text.replace('&','&amp;').replace('<','&lt;').replace('>','&gt;'),st))
        for table in doc.tables:
            data=[[cell.text for cell in row.cells] for row in table.rows]
            if data:
                t=Table(data,repeatRows=1)
                t.setStyle(TableStyle([('GRID',(0,0),(-1,-1),.4,colors.HexColor('#AAB7C4')),('BACKGROUND',(0,0),(-1,0),colors.HexColor('#17365D')),('TEXTCOLOR',(0,0),(-1,0),colors.white),('VALIGN',(0,0),(-1,-1),'TOP'),('FONTSIZE',(0,0),(-1,-1),8)]))
                story += [Spacer(1,8),t,Spacer(1,8)]
        SimpleDocTemplate(str(pdf),pagesize=A4,rightMargin=2*cm,leftMargin=2*cm,topMargin=2*cm,bottomMargin=2*cm,title=src.stem).build(story or [Paragraph(src.stem,title)])
        return pdf
    except Exception as exc:
        raise RuntimeError('No fue posible generar PDF. LibreOffice no está disponible y el fallback interno falló: '+str(exc)) from exc
