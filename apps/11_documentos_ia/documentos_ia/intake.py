from __future__ import annotations
from pathlib import Path
from .provenance import source_from_file

class IntakeError(RuntimeError): pass

def extract_text(path: str | Path) -> tuple[str, object]:
    p=Path(path)
    suffix=p.suffix.lower()
    if suffix in {'.txt','.md','.csv','.json','.html','.htm'}:
        text=p.read_text(encoding='utf-8', errors='ignore')
    elif suffix=='.docx':
        try:
            from docx import Document
        except Exception as exc:
            raise IntakeError('python-docx no disponible') from exc
        doc=Document(str(p))
        text='\n'.join(x.text for x in doc.paragraphs if x.text.strip())
        for table in doc.tables:
            for row in table.rows:
                text += '\n' + ' | '.join(c.text for c in row.cells)
    elif suffix=='.pdf':
        try:
            from pypdf import PdfReader
        except Exception as exc:
            raise IntakeError('pypdf no disponible') from exc
        reader=PdfReader(str(p))
        text='\n'.join((page.extract_text() or '') for page in reader.pages)
    else:
        raise IntakeError(f'Formato no soportado para ingesta: {suffix}')
    return text, source_from_file(p)
